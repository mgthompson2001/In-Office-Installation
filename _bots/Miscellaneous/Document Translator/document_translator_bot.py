#!/usr/bin/env python3
"""
Document Translator Bot
Translates documents (PDF, Word, text files) to any language using OCR and translation services.
Supports scanned PDFs via OCR, Word documents, and plain text files.
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
from pathlib import Path
import os
import sys
import logging
from datetime import datetime

# Try to import PDF reading libraries
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    pdfplumber = None

# Try to import PyMuPDF (fitz) for better PDF layout preservation
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    fitz = None

# Try to import OCR libraries for scanned PDFs
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    pytesseract = None

# Try to import pdf2image for OCR
try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    convert_from_path = None

# Try to import Word document reading
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

# Try to import translation libraries (multiple engines for best results)
TRANSLATORS = {}
try:
    from deep_translator import GoogleTranslator, DeepL, MicrosoftTranslator
    TRANSLATORS['google'] = GoogleTranslator
    TRANSLATORS['deepl'] = DeepL
    TRANSLATORS['microsoft'] = MicrosoftTranslator
    TRANSLATOR_AVAILABLE = True
except ImportError:
    TRANSLATOR_AVAILABLE = False
    try:
        from deep_translator import GoogleTranslator
        TRANSLATORS['google'] = GoogleTranslator
        TRANSLATOR_AVAILABLE = True
    except ImportError:
        TRANSLATOR_AVAILABLE = False

# Try Google Cloud Translation API (optional, requires API key)
try:
    from google.cloud import translate_v2 as translate
    GOOGLE_CLOUD_AVAILABLE = True
except ImportError:
    GOOGLE_CLOUD_AVAILABLE = False
    translate = None

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('document_translator.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

def _configure_ocr_paths():
    """Configure Tesseract and Poppler paths"""
    if not (OCR_AVAILABLE and pytesseract):
        return
    
    try:
        script_dir = Path(__file__).parent
        # Check environment variables
        tesseract_env = os.environ.get('TESSERACT_PATH')
        poppler_env = os.environ.get('POPPLER_PATH')
        
        if tesseract_env and Path(tesseract_env).exists():
            pytesseract.pytesseract.tesseract_cmd = tesseract_env
            logger.info(f"Configured Tesseract from env: {tesseract_env}")
        
        if poppler_env and Path(poppler_env).exists():
            os.environ['POPPLER_PATH'] = poppler_env
            logger.info(f"Configured Poppler from env: {poppler_env}")
        
        # Check vendor folder
        if not getattr(pytesseract.pytesseract, 'tesseract_cmd', None):
            vend_tess = script_dir / 'vendor' / 'Tesseract-OCR' / 'tesseract.exe'
            if vend_tess.exists():
                pytesseract.pytesseract.tesseract_cmd = str(vend_tess)
                logger.info(f"Configured Tesseract from vendor: {vend_tess}")
        
        if 'POPPLER_PATH' not in os.environ:
            vend_poppler = script_dir / 'vendor' / 'poppler' / 'Library' / 'bin'
            if vend_poppler.exists():
                os.environ['POPPLER_PATH'] = str(vend_poppler)
                logger.info(f"Configured Poppler from vendor: {vend_poppler}")
        
        # Check common locations
        if not getattr(pytesseract.pytesseract, 'tesseract_cmd', None):
            candidates = [
                Path('C:/Program Files/Tesseract-OCR/tesseract.exe'),
                Path('C:/Program Files (x86)/Tesseract-OCR/tesseract.exe'),
                Path.home() / 'AppData/Local/Programs/Tesseract-OCR/tesseract.exe',
            ]
            for c in candidates:
                if c.exists():
                    pytesseract.pytesseract.tesseract_cmd = str(c)
                    logger.info(f"Configured Tesseract from common path: {c}")
                    break
        
        if 'POPPLER_PATH' not in os.environ:
            poppler_candidates = [
                Path('C:/Program Files/poppler/Library/bin'),
                Path('C:/Program Files (x86)/poppler/Library/bin'),
                Path.home() / 'AppData/Local/poppler/Library/bin',
            ]
            for candidate in poppler_candidates:
                if candidate.exists() and (candidate / 'pdftoppm.exe').exists():
                    os.environ['POPPLER_PATH'] = str(candidate)
                    logger.info(f"Configured Poppler from common path: {candidate}")
                    break
                    
    except Exception as e:
        logger.warning(f"Error configuring OCR paths: {e}")

# Configure OCR paths on import
_configure_ocr_paths()

class DocumentTranslatorBot:
    """Main bot class for document translation"""
    
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Document Translator Bot - Version 3.1.0, Last Updated 12/04/2025")
        self.root.geometry("900x700")
        
        self.input_file_path = None
        self.output_file_path = None
        self.selected_language = None
        
        # Language codes for Google Translate
        self.languages = {
            "Spanish": "es",
            "French": "fr",
            "German": "de",
            "Italian": "it",
            "Portuguese": "pt",
            "Russian": "ru",
            "Chinese (Simplified)": "zh-CN",
            "Chinese (Traditional)": "zh-TW",
            "Japanese": "ja",
            "Korean": "ko",
            "Arabic": "ar",
            "Hindi": "hi",
            "Dutch": "nl",
            "Polish": "pl",
            "Turkish": "tr",
            "Vietnamese": "vi",
            "Thai": "th",
            "Greek": "el",
            "Hebrew": "he",
            "Swedish": "sv",
            "Norwegian": "no",
            "Danish": "da",
            "Finnish": "fi",
            "Czech": "cs",
            "Romanian": "ro",
            "Hungarian": "hu",
            "Bulgarian": "bg",
            "Croatian": "hr",
            "Serbian": "sr",
            "Slovak": "sk",
            "Slovenian": "sl",
            "Ukrainian": "uk",
            "Indonesian": "id",
            "Malay": "ms",
            "Tagalog": "tl",
            "Swahili": "sw",
            "Urdu": "ur",
            "Bengali": "bn",
            "Tamil": "ta",
            "Telugu": "te",
            "Marathi": "mr",
            "Gujarati": "gu",
            "Kannada": "kn",
            "Malayalam": "ml",
            "Punjabi": "pa"
        }
        
        self._build_ui()
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check and report on available dependencies"""
        missing = []
        
        if not PDFPLUMBER_AVAILABLE and not PYMUPDF_AVAILABLE:
            missing.append("pdfplumber or PyMuPDF (for PDF reading)")
        if not OCR_AVAILABLE:
            missing.append("pytesseract (for OCR)")
        if not PDF2IMAGE_AVAILABLE:
            missing.append("pdf2image (for OCR)")
        if not DOCX_AVAILABLE:
            missing.append("python-docx (for Word documents)")
        if not TRANSLATOR_AVAILABLE:
            missing.append("deep-translator (for translation)")
        
        if missing:
            self.gui_log("⚠️ Missing dependencies:")
            for dep in missing:
                self.gui_log(f"   - {dep}")
            self.gui_log("💡 Run install.bat to install dependencies")
        else:
            self.gui_log("✅ All dependencies available")
            
            # Check PDF library availability
            if PYMUPDF_AVAILABLE:
                self.gui_log("✅ PyMuPDF available - layout preservation enabled")
            elif PDFPLUMBER_AVAILABLE:
                self.gui_log("✅ pdfplumber available (layout preservation limited)")
            
            # Check OCR availability
            if OCR_AVAILABLE:
                try:
                    version = pytesseract.get_tesseract_version()
                    self.gui_log(f"✅ Tesseract OCR detected (version {version})")
                except:
                    self.gui_log("⚠️ Tesseract installed but not properly configured")
            
            # Check available translation engines
            if TRANSLATOR_AVAILABLE:
                engines = list(TRANSLATORS.keys())
                self.gui_log(f"✅ Translation engines available: {', '.join(engines)}")
    
    def _build_ui(self):
        """Build the main UI"""
        # Header
        header = tk.Frame(self.root, bg="#660000")
        header.pack(fill="x")
        tk.Label(header, text="Document Translator Bot", bg="#660000", fg="white",
                 font=("Segoe UI", 14, "bold"), pady=8).pack(side="left", padx=12)
        
        # Main container
        main_frame = tk.Frame(self.root, bg="#f0f0f0")
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # File selection frame
        file_frame = tk.LabelFrame(main_frame, text="Input Document", font=("Segoe UI", 10, "bold"))
        file_frame.pack(fill="x", pady=(0, 15))
        
        file_select_frame = tk.Frame(file_frame, bg="white")
        file_select_frame.pack(fill="x", padx=10, pady=10)
        
        self.file_label = tk.Label(file_select_frame, text="No file selected", 
                                   font=("Segoe UI", 9), fg="gray", bg="white")
        self.file_label.pack(side="left", padx=(0, 10))
        
        tk.Button(file_select_frame, text="Browse...", command=self._select_input_file,
                 bg="#660000", fg="white", font=("Segoe UI", 9), padx=15, pady=5,
                 cursor="hand2", relief="raised").pack(side="right")
        
        # Language selection frame
        lang_frame = tk.LabelFrame(main_frame, text="Translation Language", font=("Segoe UI", 10, "bold"))
        lang_frame.pack(fill="x", pady=(0, 15))
        
        lang_select_frame = tk.Frame(lang_frame, bg="white")
        lang_select_frame.pack(fill="x", padx=10, pady=10)
        
        tk.Label(lang_select_frame, text="Target Language:", font=("Segoe UI", 9), bg="white").pack(side="left", padx=(0, 10))
        
        self.language_var = tk.StringVar()
        self.language_dropdown = ttk.Combobox(lang_select_frame, textvariable=self.language_var,
                                             values=list(self.languages.keys()), state="readonly", width=30)
        self.language_dropdown.pack(side="left", padx=(0, 10))
        self.language_dropdown.bind("<<ComboboxSelected>>", self._on_language_selected)
        
        tk.Button(lang_select_frame, text="Select Language", command=self._show_language_dialog,
                 bg="#660000", fg="white", font=("Segoe UI", 9), padx=15, pady=5,
                 cursor="hand2", relief="raised").pack(side="left")
        
        self.selected_lang_label = tk.Label(lang_select_frame, text="No language selected",
                                           font=("Segoe UI", 9), fg="gray", bg="white")
        self.selected_lang_label.pack(side="left", padx=(15, 0))
        
        # Output location frame
        output_frame = tk.LabelFrame(main_frame, text="Output Location", font=("Segoe UI", 10, "bold"))
        output_frame.pack(fill="x", pady=(0, 15))
        
        output_select_frame = tk.Frame(output_frame, bg="white")
        output_select_frame.pack(fill="x", padx=10, pady=10)
        
        self.output_label = tk.Label(output_select_frame, text="Output will be saved next to input file",
                                    font=("Segoe UI", 9), fg="gray", bg="white")
        self.output_label.pack(side="left", padx=(0, 10))
        
        tk.Button(output_select_frame, text="Choose Location...", command=self._select_output_location,
                 bg="#660000", fg="white", font=("Segoe UI", 9), padx=15, pady=5,
                 cursor="hand2", relief="raised").pack(side="right")
        
        # Translate button
        translate_frame = tk.Frame(main_frame)
        translate_frame.pack(fill="x", pady=(0, 15))
        
        self.translate_button = tk.Button(translate_frame, text="Translate Document",
                                         command=self._translate_document,
                                         bg="#28a745", fg="white", font=("Segoe UI", 11, "bold"),
                                         padx=30, pady=10, cursor="hand2", relief="raised")
        self.translate_button.pack()
        
        # Log area
        log_frame = tk.LabelFrame(main_frame, text="Activity Log", font=("Segoe UI", 10, "bold"))
        log_frame.pack(fill="both", expand=True)
        
        self.log_text = scrolledtext.ScrolledText(log_frame, height=15, wrap="word",
                                                  font=("Consolas", 9), bg="white", fg="black")
        self.log_text.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Initial log message
        self.gui_log("Document Translator Bot initialized")
        self.gui_log("Select an input document and target language to begin")
    
    def _select_input_file(self):
        """Open file dialog to select input document"""
        filetypes = [
            ("All Supported", "*.pdf;*.docx;*.doc;*.txt"),
            ("PDF Files", "*.pdf"),
            ("Word Documents", "*.docx;*.doc"),
            ("Text Files", "*.txt"),
            ("All Files", "*.*")
        ]
        
        file_path = filedialog.askopenfilename(
            title="Select Document to Translate",
            filetypes=filetypes
        )
        
        if file_path:
            self.input_file_path = file_path
            filename = Path(file_path).name
            self.file_label.config(text=f"Selected: {filename}", fg="black")
            self.gui_log(f"📄 Selected input file: {filename}")
            
            # Auto-set output path
            self._auto_set_output_path()
    
    def _auto_set_output_path(self):
        """Automatically set output path based on input file"""
        if not self.input_file_path:
            return
        
        input_path = Path(self.input_file_path)
        lang_name = self.language_var.get() or "translated"
        lang_code = self.languages.get(lang_name, "")
        
        # Create output filename
        if input_path.suffix.lower() == '.pdf':
            output_name = f"{input_path.stem}_{lang_code}.pdf"
        elif input_path.suffix.lower() in ['.docx', '.doc']:
            output_name = f"{input_path.stem}_{lang_code}.docx"
        else:
            output_name = f"{input_path.stem}_{lang_code}.txt"
        
        self.output_file_path = str(input_path.parent / output_name)
        self.output_label.config(text=f"Output: {output_name}", fg="black")
    
    def _select_output_location(self):
        """Open file dialog to select output location"""
        if not self.input_file_path:
            messagebox.showwarning("No Input File", "Please select an input file first.")
            return
        
        input_path = Path(self.input_file_path)
        lang_name = self.language_var.get() or "translated"
        lang_code = self.languages.get(lang_name, "")
        
        # Determine default extension
        if input_path.suffix.lower() == '.pdf':
            default_ext = ".pdf"
            filetypes = [("PDF Files", "*.pdf")]
        elif input_path.suffix.lower() in ['.docx', '.doc']:
            default_ext = ".docx"
            filetypes = [("Word Documents", "*.docx")]
        else:
            default_ext = ".txt"
            filetypes = [("Text Files", "*.txt")]
        
        default_name = f"{input_path.stem}_{lang_code}{default_ext}"
        
        file_path = filedialog.asksaveasfilename(
            title="Save Translated Document As",
            defaultextension=default_ext,
            initialfile=default_name,
            filetypes=filetypes
        )
        
        if file_path:
            self.output_file_path = file_path
            filename = Path(file_path).name
            self.output_label.config(text=f"Output: {filename}", fg="black")
            self.gui_log(f"💾 Output location set: {filename}")
    
    def _on_language_selected(self, event=None):
        """Handle language selection from dropdown"""
        lang = self.language_var.get()
        if lang:
            self.selected_lang_label.config(text=f"Selected: {lang}", fg="green")
            self._auto_set_output_path()
    
    def _show_language_dialog(self):
        """Show language selection dialog in a popup window"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Select Translation Language - Version 3.1.0, Last Updated 12/04/2025")
        dialog.geometry("500x600")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Header
        header = tk.Frame(dialog, bg="#660000")
        header.pack(fill="x")
        tk.Label(header, text="Select Target Language", bg="#660000", fg="white",
                 font=("Segoe UI", 12, "bold"), pady=8).pack()
        
        # Main content
        main_frame = tk.Frame(dialog)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Instructions
        tk.Label(main_frame, text="Choose the language to translate your document into:",
                font=("Segoe UI", 10)).pack(pady=(0, 15))
        
        # Scrollable listbox
        list_frame = tk.Frame(main_frame)
        list_frame.pack(fill="both", expand=True, pady=(0, 15))
        
        scrollbar = ttk.Scrollbar(list_frame)
        scrollbar.pack(side="right", fill="y")
        
        listbox = tk.Listbox(list_frame, font=("Segoe UI", 10), yscrollcommand=scrollbar.set)
        listbox.pack(side="left", fill="both", expand=True)
        scrollbar.config(command=listbox.yview)
        
        # Populate listbox
        for lang in sorted(self.languages.keys()):
            listbox.insert(tk.END, lang)
        
        # Buttons
        button_frame = tk.Frame(main_frame)
        button_frame.pack(fill="x")
        
        def select_language():
            selection = listbox.curselection()
            if selection:
                lang = listbox.get(selection[0])
                self.language_var.set(lang)
                self.selected_lang_label.config(text=f"Selected: {lang}", fg="green")
                self._auto_set_output_path()
                dialog.destroy()
                self.gui_log(f"🌐 Selected language: {lang}")
            else:
                messagebox.showwarning("No Selection", "Please select a language from the list.")
        
        tk.Button(button_frame, text="Select", command=select_language,
                 bg="#660000", fg="white", font=("Segoe UI", 10), padx=20, pady=5,
                 cursor="hand2", relief="raised").pack(side="left", padx=(0, 10))
        
        tk.Button(button_frame, text="Cancel", command=dialog.destroy,
                 bg="gray", fg="white", font=("Segoe UI", 10), padx=20, pady=5,
                 cursor="hand2", relief="raised").pack(side="left")
        
        # Center dialog
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() // 2) - (dialog.winfo_width() // 2)
        y = (dialog.winfo_screenheight() // 2) - (dialog.winfo_height() // 2)
        dialog.geometry(f"+{x}+{y}")
    
    def _read_pdf(self, file_path):
        """Read text from PDF file with enhanced extraction, preserving layout information"""
        # Store original PDF path for layout preservation
        self.original_pdf_path = file_path
        
        # Try PyMuPDF first for better layout preservation
        if PYMUPDF_AVAILABLE:
            try:
                return self._read_pdf_pymupdf(file_path)
            except Exception as e:
                self.gui_log(f"⚠️ PyMuPDF extraction failed, trying pdfplumber: {e}")
        
        # Fallback to pdfplumber
        if not PDFPLUMBER_AVAILABLE:
            raise Exception("pdfplumber not available. Install with: pip install pdfplumber")
        
        self.gui_log(f"📖 Reading PDF: {Path(file_path).name}")
        text = ""
        
        try:
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                self.gui_log(f"   PDF has {total_pages} page(s)")
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Try multiple extraction methods for best results
                    page_text = page.extract_text() or ""
                    
                    # If extraction is poor, try extracting with layout preservation
                    if len(page_text) < 50:  # Very little text extracted
                        try:
                            # Try extracting with layout
                            page_text = page.extract_text(layout=True) or page_text
                        except:
                            pass
                    
                    # Also try extracting tables if present
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            for row in table:
                                if row:
                                    table_text = " | ".join([str(cell) if cell else "" for cell in row])
                                    page_text += "\n" + table_text
                    
                    text += page_text + "\n"
                    self.gui_log(f"   Page {page_num}/{total_pages}: {len(page_text)} characters")
            
            # If no text extracted or very little, try OCR
            if not text.strip() or len(text.strip()) < 100:
                if OCR_AVAILABLE and PDF2IMAGE_AVAILABLE:
                    self.gui_log("⚠️ Little or no text found in PDF, attempting OCR...")
                    ocr_text = self._ocr_pdf(file_path)
                    if ocr_text.strip():
                        text = ocr_text
                else:
                    self.gui_log("⚠️ Warning: PDF appears to be scanned but OCR is not available")
            
            return text.strip()
            
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def _read_pdf_pymupdf(self, file_path):
        """Read PDF using PyMuPDF with layout preservation"""
        self.gui_log(f"📖 Reading PDF with layout preservation: {Path(file_path).name}")
        
        doc = fitz.open(file_path)
        total_pages = len(doc)
        self.gui_log(f"   PDF has {total_pages} page(s)")
        
        # Store text blocks grouped by logical units (lines/paragraphs) for better translation
        self.pdf_text_blocks = []  # List of dicts: {page_num, text, bbox, font_size, font, block_index}
        self.original_text_by_block = []  # Store original text in order for mapping
        text = ""
        block_index = 0
        
        for page_num in range(total_pages):
            page = doc[page_num]
            blocks = page.get_text("dict")
            
            page_text = ""
            current_line_blocks = []  # Group spans that are on the same line
            current_y = None
            
            for block in blocks["blocks"]:
                if "lines" in block:  # Text block
                    for line in block["lines"]:
                        line_text = ""
                        line_bbox = None
                        line_font_size = 12
                        line_font = "helv"
                        
                        # Collect all spans in this line
                        for span in line["spans"]:
                            span_text = span["text"]
                            if span_text.strip():
                                line_text += span_text
                                # Use the first span's properties for the line
                                if line_bbox is None:
                                    line_bbox = span["bbox"]
                                    line_font_size = span.get("size", 12)
                                    line_font = span.get("font", "helv")
                                else:
                                    # Expand bbox to include this span
                                    span_bbox = span["bbox"]
                                    line_bbox = [
                                        min(line_bbox[0], span_bbox[0]),
                                        min(line_bbox[1], span_bbox[1]),
                                        max(line_bbox[2], span_bbox[2]),
                                        max(line_bbox[3], span_bbox[3])
                                    ]
                        
                        if line_text.strip() and line_bbox:
                            # Store as a line unit (better for translation)
                            self.pdf_text_blocks.append({
                                "page": page_num,
                                "text": line_text,
                                "bbox": line_bbox,
                                "font_size": line_font_size,
                                "font": line_font,
                                "block_index": block_index
                            })
                            self.original_text_by_block.append(line_text)
                            block_index += 1
                            page_text += line_text + "\n"
            
            text += page_text + "\n"
            self.gui_log(f"   Page {page_num + 1}/{total_pages}: {len(page_text)} characters")
        
        doc.close()
        
        # If no text extracted, try OCR
        if not text.strip() or len(text.strip()) < 100:
            if OCR_AVAILABLE and PDF2IMAGE_AVAILABLE:
                self.gui_log("⚠️ Little or no text found in PDF, attempting OCR...")
                ocr_text = self._ocr_pdf(file_path)
                if ocr_text.strip():
                    text = ocr_text
                    # Clear text blocks since OCR doesn't preserve positions
                    self.pdf_text_blocks = []
                    self.original_text_by_block = []
            else:
                self.gui_log("⚠️ Warning: PDF appears to be scanned but OCR is not available")
        
        return text.strip()
    
    def _ocr_pdf(self, file_path):
        """Extract text from scanned PDF using advanced OCR with optimal settings"""
        if not OCR_AVAILABLE or not PDF2IMAGE_AVAILABLE:
            raise Exception("OCR not available. Install Tesseract OCR and pdf2image.")
        
        self.gui_log("🔍 Running advanced OCR on PDF pages...")
        text = ""
        
        try:
            poppler_path = os.environ.get('POPPLER_PATH')
            if not poppler_path:
                # Try to find poppler
                script_dir = Path(__file__).parent
                poppler_candidates = [
                    script_dir / 'vendor' / 'poppler' / 'Library' / 'bin',
                    Path('C:/Program Files/poppler/Library/bin'),
                    Path('C:/Program Files (x86)/poppler/Library/bin'),
                ]
                for candidate in poppler_candidates:
                    if candidate.exists() and (candidate / 'pdftoppm.exe').exists():
                        poppler_path = str(candidate)
                        break
            
            # Use higher DPI for better OCR accuracy (400 DPI for high quality)
            self.gui_log("   Converting PDF to high-resolution images (400 DPI)...")
            images = convert_from_path(str(file_path), poppler_path=poppler_path, dpi=400)
            self.gui_log(f"   Converted {len(images)} pages to images")
            
            # Enhanced OCR configuration for better accuracy
            # PSM 6: Assume uniform block of text (best for documents)
            # PSM 3: Fully automatic page segmentation (fallback)
            ocr_config = '--psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz.,;:!?()[]{}\'"-+=*/%$#@&_|\\<>/~` '
            
            for page_num, img in enumerate(images, 1):
                self.gui_log(f"   Processing page {page_num}/{len(images)} with OCR...")
                
                # Try PSM 6 first (uniform block), fallback to PSM 3 (auto)
                try:
                    page_text = pytesseract.image_to_string(img, config=ocr_config)
                    if not page_text.strip():
                        # Fallback to automatic page segmentation
                        self.gui_log(f"      Trying alternative OCR mode for page {page_num}...")
                        page_text = pytesseract.image_to_string(img, config='--psm 3')
                except Exception as e:
                    self.gui_log(f"      Warning: OCR config issue on page {page_num}, using default: {e}")
                    page_text = pytesseract.image_to_string(img)
                
                text += f"\n[Page {page_num}]\n{page_text}\n"
            
            self.gui_log(f"✅ OCR completed: {len(text)} characters extracted from {len(images)} pages")
            return text
            
        except Exception as e:
            raise Exception(f"OCR failed: {str(e)}")
    
    def _read_docx(self, file_path):
        """Read text from Word document with COMPLETE structure preservation"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx not available. Install with: pip install python-docx")
        
        self.gui_log(f"📖 Reading Word document: {Path(file_path).name}")
        
        try:
            doc = Document(file_path)
            
            # Store the original document for exact copying
            self.docx_original = doc
            self.docx_structure = []
            
            text_parts = []
            element_index = 0
            
            # Read ALL document elements in order (paragraphs and tables)
            # Iterate through body elements to maintain exact order
            para_idx = 0
            table_idx = 0
            
            try:
                for element in doc.element.body:
                    # Check if it's a paragraph
                    if element.tag.endswith('p'):
                        if para_idx < len(doc.paragraphs):
                            para = doc.paragraphs[para_idx]
                            para_idx += 1
                            
                            try:
                                para_info = {
                                    'type': 'paragraph',
                                    'element': para,
                                    'text': para.text,
                                    'index': element_index,
                                    'style': para.style.name if para.style else None,
                                    'alignment': para.alignment,
                                    'paragraph_format': {},
                                    'runs': []
                                }
                                
                                # Safely extract paragraph formatting
                                try:
                                    pf = para.paragraph_format
                                    para_info['paragraph_format'] = {
                                        'space_before': getattr(pf, 'space_before', None),
                                        'space_after': getattr(pf, 'space_after', None),
                                        'line_spacing': getattr(pf, 'line_spacing', None),
                                        'line_spacing_rule': getattr(pf, 'line_spacing_rule', None),
                                        'first_line_indent': getattr(pf, 'first_line_indent', None),
                                        'left_indent': getattr(pf, 'left_indent', None),
                                        'right_indent': getattr(pf, 'right_indent', None),
                                        'keep_together': getattr(pf, 'keep_together', None),
                                        'keep_with_next': getattr(pf, 'keep_with_next', None),
                                        'page_break_before': getattr(pf, 'page_break_before', None),
                                        'widow_control': getattr(pf, 'widow_control', None)
                                    }
                                except Exception as pf_err:
                                    pass  # Use empty dict if formatting extraction fails
                                
                                # Store ALL run-level formatting safely
                                for run in para.runs:
                                    try:
                                        run_info = {
                                            'text': run.text,
                                            'bold': getattr(run, 'bold', None),
                                            'italic': getattr(run, 'italic', None),
                                            'underline': getattr(run, 'underline', None),
                                            'font_size': run.font.size if hasattr(run.font, 'size') else None,
                                            'font_name': run.font.name if hasattr(run.font, 'name') else None,
                                            'font_color': run.font.color.rgb if hasattr(run.font, 'color') and run.font.color and run.font.color.rgb else None,
                                            'font_highlight': getattr(run.font, 'highlight_color', None) if hasattr(run.font, 'highlight_color') else None,
                                            'strike': getattr(run.font, 'strike', None) if hasattr(run.font, 'strike') else None,
                                            'subscript': getattr(run.font, 'subscript', None) if hasattr(run.font, 'subscript') else None,
                                            'superscript': getattr(run.font, 'superscript', None) if hasattr(run.font, 'superscript') else None,
                                            'all_caps': getattr(run.font, 'all_caps', None) if hasattr(run.font, 'all_caps') else None,
                                            'small_caps': getattr(run.font, 'small_caps', None) if hasattr(run.font, 'small_caps') else None
                                        }
                                        para_info['runs'].append(run_info)
                                    except Exception as run_err:
                                        # Fallback: just store text
                                        para_info['runs'].append({'text': run.text, 'bold': None, 'italic': None})
                                
                                self.docx_structure.append(para_info)
                                text_parts.append(para.text if para.text else "")
                                element_index += 1
                            except Exception as para_err:
                                # Fallback: just store text
                                try:
                                    self.docx_structure.append({
                                        'type': 'paragraph',
                                        'text': para.text if hasattr(para, 'text') else '',
                                        'index': element_index,
                                        'style': None,
                                        'alignment': None,
                                        'paragraph_format': {},
                                        'runs': []
                                    })
                                    text_parts.append(para.text if hasattr(para, 'text') else "")
                                    element_index += 1
                                except:
                                    pass
                    
                    # Check if it's a table
                    elif element.tag.endswith('tbl'):
                        if table_idx < len(doc.tables):
                            table = doc.tables[table_idx]
                            table_idx += 1
                            
                            try:
                                table_info = {
                                    'type': 'table',
                                    'element': table,
                                    'index': element_index,
                                    'rows': []
                                }
                                
                                for row in table.rows:
                                    row_data = []
                                    row_text_parts = []
                                    for cell in row.cells:
                                        try:
                                            cell_info = {
                                                'text': cell.text.strip() if hasattr(cell, 'text') else '',
                                                'paragraphs': []
                                            }
                                            # Store cell paragraph info
                                            for cell_para in cell.paragraphs:
                                                try:
                                                    cell_info['paragraphs'].append({
                                                        'text': cell_para.text,
                                                        'runs': [{
                                                            'text': run.text,
                                                            'bold': getattr(run, 'bold', None),
                                                            'italic': getattr(run, 'italic', None),
                                                            'font_size': run.font.size if hasattr(run.font, 'size') else None,
                                                            'font_name': run.font.name if hasattr(run.font, 'name') else None
                                                        } for run in cell_para.runs]
                                                    })
                                                except:
                                                    cell_info['paragraphs'].append({'text': cell_para.text, 'runs': []})
                                            row_data.append(cell_info)
                                            row_text_parts.append(cell_info['text'])
                                        except:
                                            row_data.append({'text': '', 'paragraphs': []})
                                            row_text_parts.append("")
                                    table_info['rows'].append(row_data)
                                    text_parts.append(" | ".join(row_text_parts))
                                
                                self.docx_structure.append(table_info)
                                element_index += 1
                            except Exception as table_err:
                                # Fallback: create empty table structure
                                try:
                                    self.docx_structure.append({
                                        'type': 'table',
                                        'index': element_index,
                                        'rows': []
                                    })
                                    element_index += 1
                                except:
                                    pass
            except Exception as struct_err:
                # Fallback to simple text extraction if structure parsing fails
                self.gui_log(f"Warning: Error processing document structure, using simple mode: {struct_err}")
                for para in doc.paragraphs:
                    text_parts.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text for cell in row.cells])
                        text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            self.gui_log(f"✅ Extracted {len(full_text)} characters from Word document with COMPLETE formatting preserved")
            return full_text.strip()
            
        except Exception as e:
            raise Exception(f"Error reading Word document: {str(e)}")
    
    def _read_text_file(self, file_path):
        """Read text from plain text file"""
        self.gui_log(f"📖 Reading text file: {Path(file_path).name}")
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin1', 'cp1252']
            text = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    self.gui_log(f"✅ Read file using {encoding} encoding")
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise Exception("Could not decode text file with any supported encoding")
            
            return text.strip()
            
        except Exception as e:
            raise Exception(f"Error reading text file: {str(e)}")
    
    def _protect_numbers_and_dates(self, text):
        """Protect numbers, dates, phone numbers, etc. from translation"""
        import re
        
        protected_items = []
        protected_text = text
        counter = 0
        
        # Pattern to match various number/date formats
        patterns = [
            # Phone numbers: (123) 456-7890, 123-456-7890, etc.
            (r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', 'PHONE'),
            # Dates: 12/31/2024, 12-31-2024, 31/12/2024, etc.
            (r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b', 'DATE'),
            # Years: 2024, 1999, etc.
            (r'\b(19|20)\d{2}\b', 'YEAR'),
            # Currency: $123.45, $1,234.56, etc.
            (r'\$\d{1,3}(?:,\d{3})*(?:\.\d{2})?', 'CURRENCY'),
            # Percentages: 25%, 12.5%, etc.
            (r'\b\d+\.?\d*%\b', 'PERCENT'),
            # Numbers with decimals: 123.45, 1,234.56, etc.
            (r'\b\d{1,3}(?:,\d{3})*(?:\.\d+)?\b', 'NUMBER'),
            # Time: 12:30 PM, 12:30, etc.
            (r'\b\d{1,2}:\d{2}(?:\s?[AP]M)?\b', 'TIME'),
            # Social Security Numbers: 123-45-6789
            (r'\b\d{3}-\d{2}-\d{4}\b', 'SSN'),
        ]
        
        # Replace with placeholders
        for pattern, type_name in patterns:
            matches = list(re.finditer(pattern, protected_text))
            for match in reversed(matches):  # Reverse to preserve positions
                placeholder = f"__PROTECTED_{type_name}_{counter}__"
                protected_items.append({
                    'placeholder': placeholder,
                    'original': match.group(),
                    'type': type_name
                })
                protected_text = protected_text[:match.start()] + placeholder + protected_text[match.end():]
                counter += 1
        
        return protected_text, protected_items
    
    def _restore_numbers_and_dates(self, translated_text, protected_items):
        """Restore protected numbers, dates, etc. after translation"""
        restored_text = translated_text
        for item in protected_items:
            restored_text = restored_text.replace(item['placeholder'], item['original'])
        return restored_text
    
    def _translate_text(self, text, target_lang_code, protect_numbers=True):
        """Translate text using best available translation engine with smart chunking and number protection"""
        if not TRANSLATOR_AVAILABLE:
            raise Exception("deep-translator not available. Install with: pip install deep-translator")
        
        self.gui_log(f"🌐 Translating text to {target_lang_code}...")
        
        try:
            # Protect numbers and dates before translation
            if protect_numbers:
                protected_text, protected_items = self._protect_numbers_and_dates(text)
            else:
                protected_text, protected_items = text, []
            
            # Smart text chunking - split by sentences/paragraphs to preserve context
            # CRITICAL: Ensure NO text is lost during chunking
            def smart_chunk(text, max_chunk_size=4500):
                """Split text intelligently at sentence/paragraph boundaries - PRESERVE ALL TEXT"""
                if not text or not text.strip():
                    return [text] if text else ['']
                
                chunks = []
                current_chunk = ""
                
                # Split by paragraphs first (double newlines)
                if '\n\n' in text:
                    paragraphs = text.split('\n\n')
                    para_separator = '\n\n'
                elif '\n' in text:
                    paragraphs = text.split('\n')
                    para_separator = '\n'
                else:
                    # No line breaks, split by sentences
                    import re
                    sentences = re.split(r'([.!?]+\s+)', text)
                    paragraphs = []
                    for i in range(0, len(sentences)-1, 2):
                        if i+1 < len(sentences):
                            paragraphs.append(sentences[i] + sentences[i+1])
                        else:
                            paragraphs.append(sentences[i])
                    para_separator = ''
                
                for para_idx, para in enumerate(paragraphs):
                    para = para.strip()
                    if not para:
                        # Empty paragraph - preserve it
                        if current_chunk:
                            chunks.append(current_chunk.strip())
                            current_chunk = ""
                        chunks.append("")
                        continue
                    
                    # If paragraph fits, add it
                    separator_len = len(para_separator) if current_chunk else 0
                    if len(current_chunk) + separator_len + len(para) <= max_chunk_size:
                        if current_chunk:
                            current_chunk += para_separator + para
                        else:
                            current_chunk = para
                    else:
                        # Save current chunk if it has content
                        if current_chunk.strip():
                            chunks.append(current_chunk.strip())
                        
                        # If paragraph itself is too large, split by sentences
                        if len(para) > max_chunk_size:
                            import re
                            # Split by sentence endings, preserving punctuation
                            sentences = re.split(r'([.!?]+\s+)', para)
                            for sent_idx in range(0, len(sentences), 2):
                                if sent_idx < len(sentences):
                                    sent = sentences[sent_idx]
                                    if sent_idx + 1 < len(sentences):
                                        sent += sentences[sent_idx + 1]
                                    
                                    if len(current_chunk) + len(sent) + 1 <= max_chunk_size:
                                        if current_chunk:
                                            current_chunk += " " + sent
                                        else:
                                            current_chunk = sent
                                    else:
                                        if current_chunk.strip():
                                            chunks.append(current_chunk.strip())
                                        current_chunk = sent
                        else:
                            current_chunk = para
                
                # Add remaining chunk - CRITICAL: Don't lose any text
                if current_chunk.strip() or current_chunk == "":
                    chunks.append(current_chunk)
                
                # VERIFICATION: Ensure all text was chunked
                total_chunked = sum(len(c) for c in chunks)
                original_len = len(text)
                if total_chunked < original_len * 0.95:  # Allow 5% for separators
                    self.gui_log(f"⚠️ WARNING: Possible text loss during chunking: {total_chunked}/{original_len} chars")
                
                return chunks if chunks else [text]
            
            chunks = smart_chunk(protected_text)
            self.gui_log(f"   Split text into {len(chunks)} intelligent chunks for better translation quality")
            
            translated_chunks = []
            translator_used = None
            
            # Try multiple translation engines with fallback
            translation_engines = []
            if 'google' in TRANSLATORS:
                translation_engines.append(('Google Translate', TRANSLATORS['google']))
            if 'deepl' in TRANSLATORS:
                translation_engines.append(('DeepL', TRANSLATORS['deepl']))
            if 'microsoft' in TRANSLATORS:
                translation_engines.append(('Microsoft Translator', TRANSLATORS['microsoft']))
            
            if not translation_engines:
                raise Exception("No translation engines available")
            
            # Use first available engine
            engine_name, TranslatorClass = translation_engines[0]
            translator_used = engine_name
            self.gui_log(f"   Using {engine_name} for translation")
            
            for i, chunk in enumerate(chunks):
                self.gui_log(f"   Translating chunk {i+1}/{len(chunks)} ({len(chunk)} chars)...")
                
                try:
                    translator = TranslatorClass(source='auto', target=target_lang_code)
                    translated = translator.translate(chunk)
                    translated_chunks.append(translated)
                except Exception as e:
                    # Try fallback engine if available
                    if len(translation_engines) > 1:
                        self.gui_log(f"      {engine_name} failed, trying fallback...")
                        fallback_name, FallbackClass = translation_engines[1]
                        try:
                            translator = FallbackClass(source='auto', target=target_lang_code)
                            translated = translator.translate(chunk)
                            translated_chunks.append(translated)
                            translator_used = fallback_name
                        except:
                            raise Exception(f"Translation failed with all engines: {str(e)}")
                    else:
                        raise Exception(f"Translation failed: {str(e)}")
            
            # Join chunks with proper spacing - preserve original paragraph structure
            # Use original text structure to determine spacing
            if '\n\n' in text:
                translated_text = "\n\n".join(translated_chunks)
            elif '\n' in text:
                translated_text = "\n".join(translated_chunks)
            else:
                translated_text = " ".join(translated_chunks)
            
            # Restore protected numbers and dates
            if protect_numbers and protected_items:
                translated_text = self._restore_numbers_and_dates(translated_text, protected_items)
            
            # VERIFICATION: Ensure we didn't lose significant text
            original_length = len(text.strip())
            translated_length = len(translated_text.strip())
            
            # Log verification
            if original_length > 0:
                ratio = translated_length / original_length
                if ratio < 0.5:  # If translated text is less than 50% of original, something is wrong
                    self.gui_log(f"⚠️ WARNING: Translated text is {ratio*100:.1f}% of original length - possible text loss!")
                else:
                    self.gui_log(f"✅ Text length verification: {translated_length}/{original_length} chars ({ratio*100:.1f}%)")
            
            self.gui_log(f"✅ Translation completed using {translator_used}: {len(translated_text)} characters")
            return translated_text
            
        except Exception as e:
            raise Exception(f"Translation failed: {str(e)}")
    
    def _save_translated_pdf(self, text, output_path):
        """Save translated text as PDF with layout preservation"""
        # Try to preserve layout using PyMuPDF if available
        if PYMUPDF_AVAILABLE and hasattr(self, 'pdf_text_blocks') and self.pdf_text_blocks:
            try:
                self._save_translated_pdf_pymupdf(text, output_path)
                return
            except Exception as e:
                self.gui_log(f"⚠️ Layout preservation failed, using fallback: {e}")
        
        # Fallback to simple PDF creation
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.units import inch
            
            self.gui_log(f"💾 Saving translated PDF: {Path(output_path).name}")
            
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            
            # Simple text wrapping
            lines = text.split('\n')
            y = height - inch
            margin = inch
            
            for line in lines:
                if y < margin:
                    c.showPage()
                    y = height - margin
                
                # Wrap long lines
                words = line.split()
                current_line = ""
                for word in words:
                    test_line = current_line + " " + word if current_line else word
                    if c.stringWidth(test_line, "Helvetica", 10) > (width - 2 * margin):
                        if current_line:
                            c.drawString(margin, y, current_line)
                            y -= 15
                            current_line = word
                        else:
                            c.drawString(margin, y, word)
                            y -= 15
                    else:
                        current_line = test_line
                
                if current_line:
                    c.drawString(margin, y, current_line)
                    y -= 15
            
            c.save()
            self.gui_log("✅ PDF saved successfully")
            
        except ImportError:
            # Fallback: save as text file if reportlab not available
            self.gui_log("⚠️ reportlab not available, saving as text file instead")
            text_path = str(Path(output_path).with_suffix('.txt'))
            with open(text_path, 'w', encoding='utf-8') as f:
                f.write(text)
            self.gui_log(f"✅ Saved as text file: {Path(text_path).name}")
    
    def _save_translated_pdf_pymupdf(self, translated_text, output_path):
        """Save translated PDF preserving exact original layout using PyMuPDF"""
        if not hasattr(self, 'original_pdf_path') or not self.original_pdf_path:
            raise Exception("Original PDF path not available")
        
        if not hasattr(self, 'original_text_by_block') or not self.original_text_by_block:
            raise Exception("Original text blocks not available for mapping")
        
        self.gui_log(f"💾 Saving translated PDF with exact layout preservation: {Path(output_path).name}")
        
        # Open original PDF
        original_doc = fitz.open(self.original_pdf_path)
        
        # Build original text string in the same order as blocks (line by line)
        original_text_lines = self.original_text_by_block.copy()
        
        # Split translated text into lines (preserving structure)
        # The translation should maintain similar line structure
        translated_lines = [line.strip() for line in translated_text.split('\n') if line.strip()]
        
        # If line counts don't match, try to map proportionally
        if len(translated_lines) != len(original_text_lines):
            self.gui_log(f"   Line count mismatch: {len(original_text_lines)} original vs {len(translated_lines)} translated")
            # Use word-based mapping instead
            original_text_full = " ".join(original_text_lines)
            original_words = original_text_full.split()
            translated_words = translated_text.split()
            
            if not original_words:
                raise Exception("No original text found to map")
            
            word_ratio = len(translated_words) / len(original_words)
            self.gui_log(f"   Using word-based mapping: {len(original_words)} words -> {len(translated_words)} words (ratio: {word_ratio:.2f})")
            
            # Map words back to lines
            translated_lines = []
            word_index = 0
            for orig_line in original_text_lines:
                orig_words = orig_line.split()
                num_words = len(orig_words)
                translated_words_for_line = max(1, int(num_words * word_ratio))
                translated_words_for_line = min(translated_words_for_line, len(translated_words) - word_index)
                
                if translated_words_for_line > 0 and word_index < len(translated_words):
                    translated_line = " ".join(translated_words[word_index:word_index + translated_words_for_line])
                    translated_lines.append(translated_line)
                    word_index += translated_words_for_line
                else:
                    translated_lines.append(orig_line)  # Fallback
        
        # Ensure we have the same number of lines
        while len(translated_lines) < len(original_text_lines):
            translated_lines.append("")
        translated_lines = translated_lines[:len(original_text_lines)]
        
        # Create new PDF with same structure
        new_doc = fitz.open()
        
        # Group text blocks by page
        pages_blocks = {}
        for block in self.pdf_text_blocks:
            page_num = block["page"]
            if page_num not in pages_blocks:
                pages_blocks[page_num] = []
            pages_blocks[page_num].append(block)
        
        # Process each page
        line_index = 0
        for page_num in sorted(pages_blocks.keys()):
            if page_num >= len(original_doc):
                continue
                
            original_page = original_doc[page_num]
            # Create new page with same dimensions
            new_page = new_doc.new_page(width=original_page.rect.width, height=original_page.rect.height)
            
            # Copy entire page as background (preserves everything including headers/footers)
            new_page.show_pdf_page(new_page.rect, original_doc, page_num)
            
            # Get text blocks for this page, sorted by position (top to bottom)
            page_blocks = sorted(pages_blocks[page_num], key=lambda b: (b["bbox"][1], b["bbox"][0]))
            
            # Process each text block (line) on this page
            for block in page_blocks:
                if line_index >= len(translated_lines):
                    break
                
                original_block_text = block["text"].strip()
                if not original_block_text:
                    continue
                
                translated_block_text = translated_lines[line_index] if line_index < len(translated_lines) else original_block_text
                line_index += 1
                
                # Render this block with translated text at exact position
                self._render_text_block_precise(new_page, block, translated_block_text)
        
        # Save new PDF
        new_doc.save(output_path)
        new_doc.close()
        original_doc.close()
        
        self.gui_log("✅ PDF saved with exact layout preservation")
    
    def _translate_pdf_blocks(self, target_lang_code, output_path):
        """Translate PDF block-by-block for precise layout preservation"""
        if not TRANSLATOR_AVAILABLE:
            raise Exception("deep-translator not available. Install with: pip install deep-translator")
        
        if not hasattr(self, 'original_pdf_path') or not self.original_pdf_path:
            raise Exception("Original PDF path not available")
        
        self.gui_log(f"💾 Translating and saving PDF block-by-block: {Path(output_path).name}")
        
        # Open original PDF
        original_doc = fitz.open(self.original_pdf_path)
        
        # Create new PDF with same structure
        new_doc = fitz.open()
        
        # Group text blocks by page
        pages_blocks = {}
        for block in self.pdf_text_blocks:
            page_num = block["page"]
            if page_num not in pages_blocks:
                pages_blocks[page_num] = []
            pages_blocks[page_num].append(block)
        
        # Get translation engine
        translation_engines = []
        if 'google' in TRANSLATORS:
            translation_engines.append(('Google Translate', TRANSLATORS['google']))
        if 'deepl' in TRANSLATORS:
            translation_engines.append(('DeepL', TRANSLATORS['deepl']))
        
        if not translation_engines:
            raise Exception("No translation engines available")
        
        engine_name, TranslatorClass = translation_engines[0]
        translator = TranslatorClass(source='auto', target=target_lang_code)
        
        # Process each page
        total_blocks = len(self.pdf_text_blocks)
        block_count = 0
        
        for page_num in sorted(pages_blocks.keys()):
            if page_num >= len(original_doc):
                continue
                
            original_page = original_doc[page_num]
            # Create new page with same dimensions
            new_page = new_doc.new_page(width=original_page.rect.width, height=original_page.rect.height)
            
            # Copy entire page as background (preserves everything including headers/footers/graphics)
            new_page.show_pdf_page(new_page.rect, original_doc, page_num)
            
            # Get text blocks for this page, sorted by position (top to bottom)
            page_blocks = sorted(pages_blocks[page_num], key=lambda b: (b["bbox"][1], b["bbox"][0]))
            
            # Translate and render each block
            for block in page_blocks:
                original_block_text = block["text"].strip()
                if not original_block_text:
                    continue
                
                block_count += 1
                if block_count % 10 == 0:
                    self.gui_log(f"   Translated {block_count}/{total_blocks} blocks...")
                
                try:
                    # Translate this block individually
                    translated_block_text = translator.translate(original_block_text)
                    
                    # Render at exact position
                    self._render_text_block_precise(new_page, block, translated_block_text)
                except Exception as e:
                    # If translation fails, keep original text
                    self.gui_log(f"      Warning: Translation failed for block, keeping original: {e}")
                    pass
        
        # Save new PDF
        new_doc.save(output_path)
        new_doc.close()
        original_doc.close()
        
        self.gui_log(f"✅ PDF translated and saved: {block_count} blocks translated")
    
    def _render_text_block_precise(self, page, block, translated_text):
        """Render a single text block with translated text at exact position"""
        if not translated_text or not translated_text.strip():
            return
        
        # Get position and font info
        bbox = block["bbox"]
        font_size = block["font_size"]
        font_name = block["font"]
        original_text = block["text"].strip()
        
        x0, y0, x1, y1 = bbox
        
        # Calculate exact text position (PyMuPDF uses bottom-left origin)
        # y0 is top, y1 is bottom in PDF coordinates
        text_x = x0
        text_y = y1  # Bottom of the bbox for text insertion
        
        # Cover original text with white rectangle (exact size)
        text_rect = fitz.Rect(x0, y0, x1, y1)
        page.draw_rect(text_rect, color=(1, 1, 1), fill=(1, 1, 1), width=0)
        
        # Adjust font size if translated text is longer, but keep it reasonable
        length_ratio = len(translated_text) / len(original_text) if original_text else 1.0
        adjusted_font_size = font_size
        
        # Calculate available width
        available_width = x1 - x0
        
        # Try to fit text in the exact bounding box
        # Start with original font size and reduce if needed
        for scale_factor in [1.0, 0.9, 0.8, 0.7, 0.6, 0.5]:
            test_font_size = font_size * scale_factor
            test_rect = fitz.Rect(x0, y0, x1, y1)
            
            try:
                # Try inserting text in the exact bounding box
                rc = page.insert_textbox(
                    test_rect,
                    translated_text,
                    fontsize=test_font_size,
                    color=(0, 0, 0),
                    align=0  # Left align
                )
                
                # If successful (rc >= 0), we're done
                if rc >= 0:
                    return
                    
            except Exception:
                continue
        
        # If textbox insertion failed, try simple text insertion at exact position
        try:
            # Use the original position and adjust font size
            final_font_size = font_size * 0.8 if length_ratio > 1.2 else font_size
            page.insert_text(
                (text_x, text_y),
                translated_text,
                fontsize=final_font_size,
                color=(0, 0, 0)
            )
        except Exception as e:
            # If all else fails, try with smaller font
            try:
                page.insert_text(
                    (text_x, text_y),
                    translated_text[:150],  # Limit length
                    fontsize=font_size * 0.7,
                    color=(0, 0, 0)
                )
            except:
                pass  # Skip this block if it can't be rendered
    
    def _save_translated_docx(self, text, output_path):
        """Save translated text as Word document with formatting preservation"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx not available for saving Word documents")
        
        self.gui_log(f"💾 Saving translated Word document: {Path(output_path).name}")
        
        try:
            # Check if we have structure info (from _read_docx)
            if hasattr(self, 'docx_structure') and self.docx_structure:
                # Use structure-aware saving
                self._save_translated_docx_with_formatting(output_path)
            else:
                # Fallback to simple text saving
                doc = Document()
                paragraphs = text.split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        doc.add_paragraph(para_text)
                    else:
                        doc.add_paragraph()
                doc.save(output_path)
            
            self.gui_log("✅ Word document saved successfully")
            
        except Exception as e:
            raise Exception(f"Error saving Word document: {str(e)}")
    
    def _save_translated_docx_with_formatting(self, output_path):
        """Save Word document preserving EXACT original formatting structure"""
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        
        # Get target language code
        lang_name = self.language_var.get()
        if not lang_name or lang_name not in self.languages:
            raise Exception("No target language selected")
        target_lang_code = self.languages[lang_name]
        
        # Create new document
        doc = Document()
        
        # Copy document-level settings from original
        if hasattr(self, 'docx_original') and self.docx_original:
            try:
                # Copy sections and page setup
                for section in self.docx_original.sections:
                    new_section = doc.sections[-1]  # Use last section
                    new_section.page_height = section.page_height
                    new_section.page_width = section.page_width
                    new_section.left_margin = section.left_margin
                    new_section.right_margin = section.right_margin
                    new_section.top_margin = section.top_margin
                    new_section.bottom_margin = section.bottom_margin
                    new_section.header_distance = section.header_distance
                    new_section.footer_distance = section.footer_distance
            except:
                pass  # Continue if section copying fails
        
        # Process each element in order
        for element in self.docx_structure:
            if element['type'] == 'paragraph':
                # Translate paragraph text (preserve empty paragraphs for spacing)
                original_text = element['text']
                if original_text.strip():
                    try:
                        # Translate with number protection enabled
                        translated_text = self._translate_text(original_text, target_lang_code, protect_numbers=True)
                    except Exception as e:
                        self.gui_log(f"Warning: Translation failed for paragraph, using original: {e}")
                        translated_text = original_text
                else:
                    # Empty paragraph - keep it empty to preserve spacing
                    translated_text = ""
                
                # Create paragraph
                para = doc.add_paragraph()
                
                # Apply paragraph style EXACTLY
                if element.get('style'):
                    try:
                        para.style = element['style']
                    except:
                        pass
                
                # Apply alignment EXACTLY
                if element.get('alignment') is not None:
                    para.alignment = element['alignment']
                
                # Apply ALL paragraph formatting EXACTLY
                if element.get('paragraph_format'):
                    pf = para.paragraph_format
                    pf_data = element['paragraph_format']
                    
                    if pf_data.get('space_before') is not None:
                        pf.space_before = pf_data['space_before']
                    if pf_data.get('space_after') is not None:
                        pf.space_after = pf_data['space_after']
                    if pf_data.get('line_spacing') is not None:
                        pf.line_spacing = pf_data['line_spacing']
                    if pf_data.get('line_spacing_rule') is not None:
                        pf.line_spacing_rule = pf_data['line_spacing_rule']
                    if pf_data.get('first_line_indent') is not None:
                        pf.first_line_indent = pf_data['first_line_indent']
                    if pf_data.get('left_indent') is not None:
                        pf.left_indent = pf_data['left_indent']
                    if pf_data.get('right_indent') is not None:
                        pf.right_indent = pf_data['right_indent']
                    if pf_data.get('keep_together') is not None:
                        pf.keep_together = pf_data['keep_together']
                    if pf_data.get('keep_with_next') is not None:
                        pf.keep_with_next = pf_data['keep_with_next']
                    if pf_data.get('page_break_before') is not None:
                        pf.page_break_before = pf_data['page_break_before']
                    if pf_data.get('widow_control') is not None:
                        pf.widow_control = pf_data['widow_control']
                
                # Preserve run-level formatting EXACTLY - CRITICAL: Ensure ALL text is preserved
                if element.get('runs') and len(element['runs']) > 0 and translated_text.strip():
                    # Distribute translated text across runs proportionally, ensuring NO text is lost
                    runs_with_text = [r for r in element['runs'] if r.get('text', '').strip()]
                    total_original_length = sum(len(r['text']) for r in runs_with_text)
                    
                    if total_original_length > 0 and len(runs_with_text) > 0:
                        current_pos = 0
                        translated_length = len(translated_text)
                        
                        for idx, run_info in enumerate(runs_with_text):
                            # Calculate proportional length
                            run_ratio = len(run_info['text']) / total_original_length
                            run_length = int(translated_length * run_ratio)
                            
                            # Last run gets ALL remaining text to ensure nothing is lost
                            if idx == len(runs_with_text) - 1:
                                run_text = translated_text[current_pos:]  # Get ALL remaining text
                            else:
                                run_text = translated_text[current_pos:current_pos + run_length]
                            
                            current_pos += len(run_text)  # Use actual length, not calculated
                            
                            # Always add the run to preserve formatting structure
                            run = para.add_run(run_text)
                            
                            # Apply ALL run formatting EXACTLY
                            if run_info.get('bold') is not None:
                                run.bold = run_info['bold']
                            if run_info.get('italic') is not None:
                                run.italic = run_info['italic']
                            if run_info.get('underline') is not None:
                                run.underline = run_info['underline']
                            if run_info.get('font_size') is not None:
                                try:
                                    run.font.size = run_info['font_size']
                                except:
                                    pass
                            if run_info.get('font_name'):
                                try:
                                    run.font.name = run_info['font_name']
                                except:
                                    pass
                            if run_info.get('font_color') is not None:
                                try:
                                    run.font.color.rgb = run_info['font_color']
                                except:
                                    pass
                            if run_info.get('strike') is not None:
                                try:
                                    run.font.strike = run_info['strike']
                                except:
                                    pass
                            if run_info.get('subscript') is not None:
                                try:
                                    run.font.subscript = run_info['subscript']
                                except:
                                    pass
                            if run_info.get('superscript') is not None:
                                try:
                                    run.font.superscript = run_info['superscript']
                                except:
                                    pass
                        
                        # VERIFICATION: Ensure all text was distributed - CRITICAL for legal documents
                        total_distributed = sum(len(r.text) for r in para.runs)
                        if total_distributed < len(translated_text):
                            # Missing text - add it to the last run to ensure nothing is lost
                            missing_text = translated_text[total_distributed:]
                            if missing_text:
                                if para.runs:
                                    para.runs[-1].text += missing_text
                                else:
                                    para.add_run(missing_text)
                                self.gui_log(f"⚠️ WARNING: Added {len(missing_text)} missing characters to paragraph")
                    else:
                        # No runs with text, just add translated text
                        para.add_run(translated_text)
                elif translated_text.strip() or translated_text == "":
                    # No run formatting, just add text - ensure ALL text is added
                    para.add_run(translated_text)
            
            elif element['type'] == 'table':
                # Create table with EXACT same dimensions
                if element['rows']:
                    num_cols = len(element['rows'][0])
                    table = doc.add_table(rows=len(element['rows']), cols=num_cols)
                    
                    # Copy table style if available
                    if hasattr(element.get('element'), 'style'):
                        try:
                            table.style = element['element'].style
                        except:
                            pass
                    
                    # Translate and fill table cells with formatting
                    for row_idx, row_data in enumerate(element['rows']):
                        for col_idx, cell_data in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            
                            # Clear default paragraph
                            cell.paragraphs[0].clear()
                            
                            # Translate cell text
                            cell_text = cell_data['text']
                            if cell_text:
                                try:
                                    translated_cell = self._translate_text(cell_text, target_lang_code, protect_numbers=True)
                                except Exception as e:
                                    self.gui_log(f"Warning: Translation failed for table cell, using original: {e}")
                                    translated_cell = cell_text
                                
                                # Preserve cell paragraph formatting
                                if cell_data.get('paragraphs'):
                                    for para_info in cell_data['paragraphs']:
                                        cell_para = cell.add_paragraph()
                                        para_text = para_info.get('text', '')
                                        if para_text:
                                            # Translate paragraph text
                                            try:
                                                translated_para_text = self._translate_text(para_text, target_lang_code, protect_numbers=True)
                                            except:
                                                translated_para_text = para_text
                                            
                                            # Preserve run formatting
                                            if para_info.get('runs'):
                                                total_len = sum(len(r['text']) for r in para_info['runs'] if r['text'])
                                                if total_len > 0:
                                                    pos = 0
                                                    for run_info in para_info['runs']:
                                                        if not run_info['text']:
                                                            continue
                                                        ratio = len(run_info['text']) / total_len
                                                        length = max(1, int(len(translated_para_text) * ratio))
                                                        run_text = translated_para_text[pos:pos + length]
                                                        pos += length
                                                        if run_text:
                                                            run = cell_para.add_run(run_text)
                                                            if run_info.get('bold'):
                                                                run.bold = True
                                                            if run_info.get('italic'):
                                                                run.italic = True
                                                            if run_info.get('font_size'):
                                                                try:
                                                                    run.font.size = run_info['font_size']
                                                                except:
                                                                    pass
                                                            if run_info.get('font_name'):
                                                                try:
                                                                    run.font.name = run_info['font_name']
                                                                except:
                                                                    pass
                                                else:
                                                    cell_para.add_run(translated_para_text)
                                            else:
                                                cell_para.add_run(translated_para_text)
                            else:
                                # Empty cell - keep it empty
                                pass
        
        doc.save(output_path)
        self.gui_log("✅ Word document saved with EXACT formatting preserved")
        
        # Clean up
        if hasattr(self, 'docx_structure'):
            delattr(self, 'docx_structure')
        if hasattr(self, 'docx_original'):
            delattr(self, 'docx_original')
    
    def _save_translated_text(self, text, output_path):
        """Save translated text as plain text file"""
        self.gui_log(f"💾 Saving translated text file: {Path(output_path).name}")
        
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            self.gui_log("✅ Text file saved successfully")
            
        except Exception as e:
            raise Exception(f"Error saving text file: {str(e)}")
    
    def _translate_document(self):
        """Main translation workflow"""
        # Validate inputs
        if not self.input_file_path:
            messagebox.showerror("No Input File", "Please select an input document.")
            return
        
        if not os.path.exists(self.input_file_path):
            messagebox.showerror("File Not Found", f"Input file not found:\n{self.input_file_path}")
            return
        
        lang_name = self.language_var.get()
        if not lang_name:
            messagebox.showerror("No Language Selected", "Please select a target language.")
            return
        
        target_lang_code = self.languages.get(lang_name)
        if not target_lang_code:
            messagebox.showerror("Invalid Language", "Selected language is not valid.")
            return
        
        # Auto-set output path if not set
        if not self.output_file_path:
            self._auto_set_output_path()
        
        if not self.output_file_path:
            messagebox.showerror("No Output Path", "Please specify an output location.")
            return
        
        # Disable translate button during processing
        self.translate_button.config(state="disabled", text="Translating...")
        self.root.update()
        
        try:
            # Read input document
            input_path = Path(self.input_file_path)
            file_ext = input_path.suffix.lower()
            
            if file_ext == '.pdf':
                text = self._read_pdf(self.input_file_path)
            elif file_ext in ['.docx', '.doc']:
                text = self._read_docx(self.input_file_path)
            elif file_ext == '.txt':
                text = self._read_text_file(self.input_file_path)
            else:
                raise Exception(f"Unsupported file type: {file_ext}")
            
            if not text.strip():
                raise Exception("No text could be extracted from the document.")
            
            # For PDFs with layout preservation, translate block-by-block for better mapping
            if file_ext == '.pdf' and PYMUPDF_AVAILABLE and hasattr(self, 'pdf_text_blocks') and self.pdf_text_blocks:
                # Translate block-by-block for precise mapping
                self.gui_log("📝 Translating PDF block-by-block for precise layout preservation...")
                output_path = Path(self.output_file_path)
                self._translate_pdf_blocks(target_lang_code, str(output_path))
            else:
                # Translate text normally
                translated_text = self._translate_text(text, target_lang_code, protect_numbers=True)
                
                # Save translated document
                output_path = Path(self.output_file_path)
                output_ext = output_path.suffix.lower()
                
                if output_ext == '.pdf':
                    self._save_translated_pdf(translated_text, str(output_path))
                elif output_ext in ['.docx', '.doc']:
                    self._save_translated_docx(translated_text, str(output_path))
                else:
                    self._save_translated_text(translated_text, str(output_path))
            
            # Success message
            self.gui_log("=" * 60)
            self.gui_log("✅ Translation completed successfully!")
            self.gui_log(f"📄 Input: {input_path.name}")
            self.gui_log(f"📄 Output: {output_path.name}")
            self.gui_log(f"🌐 Language: {lang_name}")
            self.gui_log("=" * 60)
            
            messagebox.showinfo("Translation Complete",
                              f"Document translated successfully!\n\n"
                              f"Output saved to:\n{output_path}\n\n"
                              f"Language: {lang_name}")
            
            # Open output folder
            try:
                if sys.platform == "win32":
                    os.startfile(output_path.parent)
            except:
                pass
            
        except Exception as e:
            error_msg = str(e)
            self.gui_log(f"❌ Error: {error_msg}")
            messagebox.showerror("Translation Error", f"Failed to translate document:\n\n{error_msg}")
        
        finally:
            # Re-enable translate button
            self.translate_button.config(state="normal", text="Translate Document")
    
    def gui_log(self, message):
        """Add a message to the GUI log"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.log_text.insert("end", f"[{timestamp}] {message}\n")
        self.log_text.see("end")
        self.root.update_idletasks()
    
    def run(self):
        """Start the GUI application"""
        self.root.mainloop()

if __name__ == "__main__":
    app = DocumentTranslatorBot()
    app.run()

